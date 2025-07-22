import hashlib
import json
from pathlib import Path
from typing import Union, Dict, Any
from docx import Document


def hash_document_and_json(
    doc_path: Union[str, Path],
    json_data: Union[Dict[str, Any], str, Path],
    doc_type: str = "pdf"
) -> str:
    """Hash a document (PDF or Word) together with JSON data.
    
    Args:
        doc_path: Path to the document file (PDF or Word)
        json_data: Either a dictionary of JSON data, a JSON string, or a path to a 
                  JSON file
        doc_type: Type of document ("pdf" or "word")
        
    Returns:
        A SHA-256 hash of the combined document and JSON data
    """
    # Initialize SHA-256 hash object
    hash_obj = hashlib.sha256()
    
    # Hash the document file
    doc_path = Path(doc_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"Document file not found: {doc_path}")
    
    if doc_type.lower() == "pdf":
        # Hash PDF file in chunks
        with open(doc_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hash_obj.update(chunk)
    elif doc_type.lower() == "word":
        # Hash Word file content
        try:
            doc = Document(doc_path)
            # Extract text content
            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            # Join all text and hash it
            doc_text = "\n".join(text_content)
            doc_bytes = doc_text.encode('utf-8')
            hash_obj.update(doc_bytes)
        except Exception as e:
            raise ValueError(f"Error processing Word file: {str(e)}")
    else:
        raise ValueError(f"Unsupported document type: {doc_type}")
    
    # Process JSON data
    if isinstance(json_data, (str, Path)):
        # If it's a string path, read the JSON file
        json_path = Path(json_data)
        if json_path.exists():
            with open(json_path, 'r', encoding='utf-8') as f:
                json_str = f.read()
        else:
            # Assume it's a JSON string
            json_str = str(json_data)
    else:
        # Convert dictionary to sorted JSON string for consistent hashing
        json_str = json.dumps(json_data, sort_keys=True)
    
    # Hash the JSON data - ensure it's bytes before updating
    json_bytes = json_str.encode('utf-8')
    hash_obj.update(json_bytes)
    
    return hash_obj.hexdigest()


def verify_hash(
    doc_path: Union[str, Path],
    json_data: Union[Dict[str, Any], str, Path],
    expected_hash: str,
    doc_type: str = "pdf"
) -> bool:
    """Verify if the current hash of document and JSON matches the expected hash.
    
    Args:
        doc_path: Path to the document file (PDF or Word)
        json_data: Either a dictionary of JSON data, a JSON string, or a path to a 
                  JSON file
        expected_hash: The expected hash to compare against
        doc_type: Type of document ("pdf" or "word")
        
    Returns:
        True if the hash matches, False otherwise
    """
    current_hash = hash_document_and_json(doc_path, json_data, doc_type)
    return current_hash == expected_hash 